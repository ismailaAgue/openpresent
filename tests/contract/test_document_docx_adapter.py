"""Contract tests for DocumentDocxExportAdapter (ADR-041, v3 Phase 3)."""

import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.models.recipe import Recipe, Outline, Slide, ContentBlock, BlockType, StructureSource, Theme
from backend.adapters.export.document_docx_adapter import DocumentDocxExportAdapter


def make_recipe(slides=None):
    outline = Outline(structure_source=StructureSource.AI_GENERATED, slides=slides or [
        Slide(order=1, title="AI-First Market Report", content_blocks=[]),
        Slide(order=2, title="Executive Summary", content_blocks=[
            ContentBlock(
                type=BlockType.BULLET,
                text="The market for AI-first tools grew significantly this year.",
            ),
        ]),
        Slide(order=3, title="Key Findings", content_blocks=[
            ContentBlock(type=BlockType.BULLET, text="Adoption is accelerating"),
            ContentBlock(type=BlockType.BULLET, text="Pricing pressure is increasing"),
            ContentBlock(type=BlockType.NOTE, text="Mention this only if asked."),
        ]),
    ])
    return Recipe.new(project_id="p1", source_text="Topic: test", outline=outline,
                       theme=Theme(), audience_type="general", language="en")


def test_format_id():
    assert DocumentDocxExportAdapter().format_id() == "document_docx"


def test_produces_a_valid_docx():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))  # raises if not a real docx
    assert len(doc.paragraphs) > 0


def test_first_slide_becomes_the_title_not_a_numbered_slide_heading():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI-First Market Report" in full_text
    assert "Slide 1" not in full_text  # this is a document, not slide-notes framing


def test_section_headings_present_for_every_non_title_slide():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Executive Summary" in heading_texts
    assert "Key Findings" in heading_texts


def test_single_sentence_bullet_runs_as_a_paragraph_not_a_list_item():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    # The Executive Summary slide has exactly one bullet that reads as a
    # full sentence -> should be a plain paragraph, not "List Bullet" style.
    matching = [p for p in doc.paragraphs if "grew significantly" in p.text]
    assert len(matching) == 1
    assert matching[0].style.name != "List Bullet"


def test_multiple_bullets_stay_a_real_bullet_list():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    bullet_paragraphs = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = [p.text for p in bullet_paragraphs]
    assert "Adoption is accelerating" in bullet_texts
    assert "Pricing pressure is increasing" in bullet_texts


def test_speaker_notes_are_not_rendered_in_the_document():
    """Notes exist for a presenter reading a deck aloud — irrelevant
    once the deliverable IS the document, unlike the notes_docx
    adapter where they're the entire point."""
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Mention this only if asked." not in full_text


def test_empty_outline_still_produces_a_valid_docx():
    empty_outline = Outline(structure_source=StructureSource.AI_GENERATED, slides=[])
    empty_recipe = Recipe.new(project_id="p1", source_text="Topic: test", outline=empty_outline,
                               theme=Theme(), audience_type="general", language="en")
    output = DocumentDocxExportAdapter().export(empty_recipe)
    doc = Document(io.BytesIO(output))  # must not raise
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Untitled Document" in full_text


# -- Multi-paragraph prose and theme design (ADR-054) -----------------------

def test_multiple_paragraph_length_bullets_become_multiple_real_paragraphs():
    """The AI content prompt, for document_docx, now asks for 1-3
    SEPARATE paragraph-bullets per section (not one long joined bullet)
    — each one must render as its own real paragraph, not a bulleted
    list of paragraph-length chunks (which would still look deck-ish,
    undermining the whole point of asking for real prose)."""
    slides = [
        Slide(order=1, title="Report", content_blocks=[]),
        Slide(order=2, title="Market Overview", content_blocks=[
            ContentBlock(type=BlockType.BULLET,
                         text="The market grew substantially this year, driven by strong demand."),
            ContentBlock(type=BlockType.BULLET,
                         text="Competition also intensified, with several new entrants."),
        ]),
    ]
    output = DocumentDocxExportAdapter().export(make_recipe(slides=slides))
    doc = Document(io.BytesIO(output))
    matching = [p for p in doc.paragraphs if p.text.strip() and p.text.strip() not in ("Report", "Market Overview")]
    bullet_style_paragraphs = [p for p in matching if p.style.name == "List Bullet"]
    assert len(bullet_style_paragraphs) == 0  # neither paragraph got a bullet-list style
    assert any("grew substantially" in p.text for p in matching)
    assert any("intensified" in p.text for p in matching)


def test_heading_color_matches_the_projects_theme():
    """Before ADR-054, this adapter never used the project's theme at
    all — every document rendered in Word's default blue heading
    color regardless of the chosen theme, unlike every other export
    format. Section headings must now carry the theme's real title
    color."""
    from backend.adapters.export.pptx_adapter import _COLOR_SETS
    from docx.shared import RGBColor

    recipe = make_recipe()
    recipe.theme.color_set_id = "warm_editorial"
    output = DocumentDocxExportAdapter().export(recipe)
    doc = Document(io.BytesIO(output))

    heading_paragraphs = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(heading_paragraphs) > 0
    expected_color = RGBColor(*_COLOR_SETS["warm_editorial"]["title"])
    assert heading_paragraphs[0].runs[0].font.color.rgb == expected_color


def test_different_themes_produce_different_heading_colors():
    output_a = DocumentDocxExportAdapter().export(_recipe_with_theme("neutral"))
    output_b = DocumentDocxExportAdapter().export(_recipe_with_theme("modern_dark"))
    doc_a = Document(io.BytesIO(output_a))
    doc_b = Document(io.BytesIO(output_b))
    color_a = next(p for p in doc_a.paragraphs if p.style.name.startswith("Heading")).runs[0].font.color.rgb
    color_b = next(p for p in doc_b.paragraphs if p.style.name.startswith("Heading")).runs[0].font.color.rgb
    assert color_a != color_b


def _recipe_with_theme(color_set_id: str) -> Recipe:
    recipe = make_recipe()
    recipe.theme.color_set_id = color_set_id
    return recipe


def test_title_page_has_a_real_full_width_accent_rule():
    """ADR-054: the title page no longer relies on python-docx's
    built-in 'Title' style (whose border box doesn't track the actual
    centered text width — confirmed by rendering and looking, see the
    module docstring). Instead a dedicated paragraph carries a real
    OOXML bottom border. Checked at the XML level since python-docx's
    high-level API has no paragraph-border getter."""
    from docx.oxml.ns import qn

    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    found_border = False
    for p in doc.paragraphs[:5]:  # the rule is one of the first few paragraphs, on the title page
        p_pr = p._p.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:pBdr")) is not None:
            found_border = True
            break
    assert found_border


def test_title_text_is_centered():
    output = DocumentDocxExportAdapter().export(make_recipe())
    doc = Document(io.BytesIO(output))
    title_para = next(p for p in doc.paragraphs if "AI-First Market Report" in p.text)
    assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_title_color_matches_theme():
    from backend.adapters.export.pptx_adapter import _COLOR_SETS
    from docx.shared import RGBColor

    recipe = _recipe_with_theme("blue_academic")
    output = DocumentDocxExportAdapter().export(recipe)
    doc = Document(io.BytesIO(output))
    title_para = next(p for p in doc.paragraphs if "AI-First Market Report" in p.text)
    expected_color = RGBColor(*_COLOR_SETS["blue_academic"]["title"])
    assert title_para.runs[0].font.color.rgb == expected_color

