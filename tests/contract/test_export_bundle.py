import io
import zipfile
from docx import Document
from backend.models.recipe import Recipe, Outline, Slide, ContentBlock, BlockType, StructureSource, Theme
from backend.adapters.export.docx_notes_adapter import SpeakerNotesDocxExportAdapter
from backend.engines.export_bundle import build_export_bundle


def make_recipe():
    outline = Outline(structure_source=StructureSource.AI_GENERATED, slides=[
        Slide(order=1, title="Intro", content_blocks=[
            ContentBlock(type=BlockType.BULLET, text="Welcome"),
            ContentBlock(type=BlockType.NOTE, text="Say hello warmly."),
        ]),
        Slide(order=2, title="No Notes Slide", content_blocks=[
            ContentBlock(type=BlockType.BULLET, text="Just a bullet"),
        ]),
    ])
    return Recipe.new(project_id="p1", source_text="Topic: test", outline=outline,
                       theme=Theme(), audience_type="general", language="en")


def test_docx_notes_adapter_format_id():
    assert SpeakerNotesDocxExportAdapter().format_id() == "notes_docx"


def test_docx_notes_adapter_produces_valid_docx_with_titles_and_notes():
    recipe = make_recipe()
    output = SpeakerNotesDocxExportAdapter().export(recipe)
    doc = Document(io.BytesIO(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Intro" in full_text
    assert "Say hello warmly." in full_text
    assert "No Notes Slide" in full_text
    assert "No speaker notes" in full_text  # second slide has no NOTE block


def test_export_bundle_contains_both_files():
    recipe = make_recipe()
    fake_pptx_bytes = b"FAKE_PPTX_CONTENT"
    zip_bytes = build_export_bundle(recipe, fake_pptx_bytes, "pptx")

    zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
    names = zf.namelist()
    assert "presentation.pptx" in names
    assert "speaker_notes.docx" in names
    assert zf.read("presentation.pptx") == fake_pptx_bytes

    # speaker_notes.docx is a real, parseable docx, not just bytes
    notes_doc = Document(io.BytesIO(zf.read("speaker_notes.docx")))
    assert any("Intro" in p.text for p in notes_doc.paragraphs)
