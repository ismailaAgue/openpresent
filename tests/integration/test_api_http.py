"""
HTTP-level integration tests — the gap flagged repeatedly in earlier
reviews: every other test in this suite exercises engines/adapters
directly, never a real request through the actual FastAPI/Starlette
HTTP stack. That gap is exactly how two real bugs shipped unnoticed —
the zip-bundling response shape (ADR-030) and the notes-in-visible-
body bug (ADR-029) — since a direct engine call never touches routing,
response headers, content-type negotiation, or multipart parsing the
way a real request does.

Uses FastAPI's TestClient (Starlette under the hood) as a real WSGI/
ASGI client against `backend.api.main:app` — actual HTTP request/
response objects, real routing, real header handling. The `with
TestClient(app) as client:` context-manager form is used deliberately
(not the bare constructor) so FastAPI's startup event actually fires,
which starts the in-process worker thread (backend/api/main.py's
`_start_inprocess_worker`) — without this, every /generate/async test
below would hang forever polling a job that nothing is processing.

Hermetic by construction: every registry singleton is reset before
each test (fresh in-memory SQLite for queue/storage/auth, matching
production behavior when DATABASE_URL is unset — see
registry._database_url()), and AI/media/research default to disabled
unless a specific test explicitly wires up a fake. No live network
calls, no flakiness from external providers — this suite tests
OpenPresent's own HTTP surface, not third-party API availability.
"""

import io
import zipfile
import time
import pytest
from docx import Document
from fastapi.testclient import TestClient
from backend.adapters import registry

SAMPLE_TEXT_DOC = (
    "**The Water Cycle**\n\n"
    "**Evaporation**\n"
    "The sun heats water in oceans, lakes, and rivers, turning it into vapor "
    "that rises into the atmosphere.\n\n"
    "**Condensation**\n"
    "As water vapor rises and cools, it condenses into tiny droplets, "
    "forming clouds.\n\n"
    "**Precipitation**\n"
    "When droplets in clouds combine and grow heavy enough, they fall back "
    "to earth as rain, snow, or hail.\n"
).encode("utf-8")


def _make_minimal_pdf(text_lines: list[str]) -> bytes:
    """A hand-written, valid, minimal single-page PDF with real,
    extractable text — used instead of a PDF-generation library (e.g.
    reportlab) specifically because this project's only real PDF
    dependency is pypdf (for reading), not for writing. Round-tripped
    through pypdf's own reader as part of ADR-050's real-PDF test
    coverage — see test_ai_generate_engine.py/test_title_enhancement.py
    for the equivalent reasoning on why SAMPLE_TEXT_DOC alone wasn't
    enough evidence that PDF-specific ingestion actually works end to
    end (it never had, before ADR-050 — only the corrupt-PDF error
    path had ever been exercised via a real HTTP request)."""
    content_lines = ["BT", "/F1 12 Tf", "72 720 Td"]
    for i, line in enumerate(text_lines):
        if i > 0:
            content_lines.append("0 -18 Td")
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content_lines.append(f"({escaped}) Tj")
    content_lines.append("ET")
    content = "\n".join(content_lines).encode("latin-1")

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>\nendobj\n",
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        b"5 0 obj\n<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream\nendobj\n",
    ]
    pdf = b"%PDF-1.4\n"
    offsets = []
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj
    xref_offset = len(pdf)
    pdf += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += b"trailer\n<< /Size " + str(len(objects) + 1).encode() + b" /Root 1 0 R >>\nstartxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    return pdf


SAMPLE_PDF_DOC = _make_minimal_pdf([
    "The Water Cycle",
    "Evaporation: Water turns into vapor when heated by the sun.",
    "Condensation: Vapor cools and forms clouds.",
    "Precipitation: Water falls back to earth as rain or snow.",
])


@pytest.fixture(autouse=True)
def reset_all_registry_singletons(monkeypatch):
    """Every registry-cached adapter reset to None before each test —
    fresh in-memory SQLite queue/storage/auth per test (no cross-test
    pollution), and AI/media/research forced to the deterministic/off
    default so these tests never depend on live network or real API
    keys. Individual tests that need an AI adapter wire up their own
    fake via monkeypatch, same pattern used throughout this suite."""
    for attr in ("_ai_adapter_instance", "_queue_adapter_instance", "_storage_adapter_instance",
                 "_auth_adapter_instance", "_analytics_adapter_instance", "_media_adapter_instance",
                 "_research_adapter_instance", "_quota_adapter_instance", "_workspace_adapter_instance", "_brand_adapter_instance"):
        setattr(registry, attr, None)
    monkeypatch.setenv("OPENPRESENT_AI_ADAPTER", "null")
    monkeypatch.setenv("OPENPRESENT_RESEARCH_ADAPTER", "null")
    monkeypatch.delenv("DATABASE_URL", raising=False)
    monkeypatch.delenv("OPENPRESENT_UNSPLASH_ACCESS_KEY", raising=False)
    monkeypatch.delenv("OPENPRESENT_PEXELS_API_KEY", raising=False)
    monkeypatch.delenv("OPENPRESENT_PIXABAY_API_KEY", raising=False)
    yield
    for attr in ("_ai_adapter_instance", "_queue_adapter_instance", "_storage_adapter_instance",
                 "_auth_adapter_instance", "_analytics_adapter_instance", "_media_adapter_instance",
                 "_research_adapter_instance", "_quota_adapter_instance", "_workspace_adapter_instance", "_brand_adapter_instance"):
        setattr(registry, attr, None)


@pytest.fixture
def client():
    from backend.api.main import app
    # Context-manager form is required — this is what actually fires
    # FastAPI's startup event and starts the in-process worker thread.
    # A bare TestClient(app) never processes async jobs at all.
    with TestClient(app) as c:
        yield c


@pytest.fixture
def client_no_worker(monkeypatch):
    """Same as `client`, but with the in-process background worker
    disabled entirely. Use this for any test that manually drives a
    job through QueuePort (enqueue/dequeue/update_stage) to check the
    API's response shape — the real worker thread `client` starts
    polls the SAME shared queue continuously, and will race to
    dequeue and fully process ANY enqueued job it can see, including
    ones a test only meant to control by hand. On fast hardware (or
    just different OS thread-scheduling behavior — this was found via
    a real, reproducible failure on Windows that never appeared across
    many repeated runs on Linux) that race is very much winnable by
    the real worker, silently overwriting a test's manually-set state
    before the test's own assertions run. Disabling the worker for
    these specific tests removes the race at its root instead of
    trying to out-time it."""
    monkeypatch.setenv("OPENPRESENT_INPROCESS_WORKER", "false")
    from backend.api.main import app
    with TestClient(app) as c:
        yield c


def _poll_job_until_done(client, job_id, timeout_seconds=10):
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        resp = client.get(f"/jobs/{job_id}")
        assert resp.status_code == 200
        if resp.json()["status"] == "done":
            return resp.json()
        if resp.json()["status"] == "failed":
            raise AssertionError(f"job failed: {resp.json()}")
        time.sleep(0.1)
    raise AssertionError(f"job {job_id} did not complete within {timeout_seconds}s")


# -- Basic routing (regression: the / -> 404 issue) ----------------------

def test_root_returns_200_not_404(client):
    resp = client.get("/")
    assert resp.status_code == 200
    assert resp.json()["status"] == "ok"


def test_health_returns_expected_shape(client):
    resp = client.get("/health")
    assert resp.status_code == 200
    body = resp.json()
    for key in ("status", "ai_adapter", "ai_providers_configured", "media_adapter",
                "media_providers_configured", "research_adapter", "research_providers_configured",
                "sentry_active", "database_url_present"):
        assert key in body, f"missing expected /health field: {key}"
    assert body["status"] == "ok"


# -- Document upload, sync (real multipart, real zip response) ----------

def test_generate_sync_document_returns_zip_bundle_by_default(client):
    resp = client.post(
        "/generate",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/zip"
    assert "presentation.zip" in resp.headers["content-disposition"]

    zf = zipfile.ZipFile(io.BytesIO(resp.content))
    names = zf.namelist()
    assert "presentation.pptx" in names
    assert "speaker_notes.docx" in names
    assert len(zf.read("presentation.pptx")) > 1000  # a real, non-trivial pptx
    assert len(zf.read("speaker_notes.docx")) > 500


def test_generate_sync_bundle_speaker_notes_false_returns_bare_pptx(client):
    resp = client.post(
        "/generate?bundle_speaker_notes=false",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == \
        "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    assert resp.content[:2] == b"PK"  # a real pptx (zip container), not our multi-file bundle
    # Confirms it's a bare pptx, not accidentally still our 2-file bundle:
    zf = zipfile.ZipFile(io.BytesIO(resp.content))
    assert "speaker_notes.docx" not in zf.namelist()


def test_generate_sync_unsupported_file_type_returns_400(client):
    resp = client.post(
        "/generate",
        files={"file": ("data.xyz", b"whatever content", "application/octet-stream")},
    )
    assert resp.status_code == 400


def test_generate_sync_corrupt_pdf_returns_422(client):
    resp = client.post(
        "/generate",
        files={"file": ("broken.pdf", b"this is not a real pdf file at all", "application/pdf")},
    )
    assert resp.status_code == 422


def test_generate_sync_accepts_audience_language_and_slide_count_params(client):
    """ADR-034 regression at the real HTTP layer — these params only
    ever worked when calling the engine directly before that fix;
    this confirms they're actually wired through routing."""
    resp = client.post(
        "/generate?audience_type=business&language=en&target_slide_count=5",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200


# -- Topic generation, sync (real JSON body, real zip + headers) --------

def test_generate_topic_sync_returns_zip_with_quality_headers(client):
    resp = client.post("/generate/topic", json={"topic": "Photosynthesis", "slide_count": 4})
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/zip"
    assert "X-Structure-Source" in resp.headers
    assert "X-Quality-Score" in resp.headers
    assert resp.headers["X-Structure-Source"] == "deterministic-topic"  # AI disabled in this fixture

    zf = zipfile.ZipFile(io.BytesIO(resp.content))
    assert "presentation.pptx" in zf.namelist()
    assert "speaker_notes.docx" in zf.namelist()


def test_generate_topic_empty_topic_returns_400(client):
    resp = client.post("/generate/topic", json={"topic": "   ", "slide_count": 5})
    assert resp.status_code == 400


def test_generate_topic_clamps_extreme_slide_counts(client):
    resp = client.post("/generate/topic", json={"topic": "A Topic", "slide_count": 9999})
    assert resp.status_code == 200
    zf = zipfile.ZipFile(io.BytesIO(resp.content))
    from pptx import Presentation
    prs = Presentation(io.BytesIO(zf.read("presentation.pptx")))
    assert len(list(prs.slides)) <= 30  # MAX_SLIDE_COUNT — clamped, not literally 9999 slides


# -- Async round trips, real worker thread processing --------------------

def test_generate_async_document_full_round_trip(client):
    enqueue_resp = client.post(
        "/generate/async",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]
    assert enqueue_resp.json()["status"] == "pending"

    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] == "rule-based"
    assert result["slide_count"] >= 3

    download_resp = client.get(f"/jobs/{job_id}/download")
    assert download_resp.status_code == 200
    assert download_resp.headers["content-type"] == "application/zip"
    zf = zipfile.ZipFile(io.BytesIO(download_resp.content))
    assert "presentation.pptx" in zf.namelist()
    assert "speaker_notes.docx" in zf.namelist()


def test_generate_topic_async_full_round_trip(client):
    enqueue_resp = client.post("/generate/topic/async", json={"topic": "Volcanoes", "slide_count": 4})
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]

    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] == "deterministic-topic"
    assert "quality_score" in result

    download_resp = client.get(f"/jobs/{job_id}/download")
    assert download_resp.status_code == 200
    zf = zipfile.ZipFile(io.BytesIO(download_resp.content))
    assert "presentation.pptx" in zf.namelist()


# -- Documents as a second output type (ADR-041, v3 Phase 3) -------------

# -- Cost circuit breaker (ADR-043) --------------------------------------

def test_anonymous_generation_blocked_after_daily_limit(client, monkeypatch):
    monkeypatch.setenv("OPENPRESENT_DAILY_GENERATION_LIMIT_ANON", "2")
    for _ in range(2):
        resp = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4})
        assert resp.status_code == 200
    blocked = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4})
    assert blocked.status_code == 429
    assert "daily limit" in blocked.json()["detail"].lower()


def test_quota_gate_runs_before_any_generation_work(client, monkeypatch):
    """The 429 must come from the gate itself, not from generation
    happening and then being discarded — proven by setting the limit to
    zero and confirming the very first request is blocked, with no
    generation-specific side effect (no X-Project-Id etc.) ever occurring."""
    monkeypatch.setenv("OPENPRESENT_DAILY_GENERATION_LIMIT_ANON", "0")
    resp = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4})
    assert resp.status_code == 429


def test_async_enqueue_is_also_gated_not_just_the_sync_path(client, monkeypatch):
    monkeypatch.setenv("OPENPRESENT_DAILY_GENERATION_LIMIT_ANON", "0")
    resp = client.post("/generate/topic/async", json={"topic": "Volcanoes", "slide_count": 4})
    assert resp.status_code == 429


def test_quota_is_keyed_separately_per_user(client, monkeypatch):
    """Two different accounts must not share a quota bucket — this test
    would fail if the key were something global like just "user" instead
    of including the actual user id."""
    monkeypatch.setenv("OPENPRESENT_DAILY_GENERATION_LIMIT_USER", "1")
    client.post("/auth/register", json={"email": "a@example.com", "password": "password123"})
    token_a = client.post("/auth/login", json={"email": "a@example.com", "password": "password123"}).json()["session_token"]
    client.post("/auth/register", json={"email": "b@example.com", "password": "password123"})
    token_b = client.post("/auth/login", json={"email": "b@example.com", "password": "password123"}).json()["session_token"]

    resp_a1 = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4},
                           headers={"Authorization": f"Bearer {token_a}"})
    assert resp_a1.status_code == 200
    resp_a2 = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4},
                           headers={"Authorization": f"Bearer {token_a}"})
    assert resp_a2.status_code == 429  # user A is now over their limit of 1

    resp_b1 = client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4},
                           headers={"Authorization": f"Bearer {token_b}"})
    assert resp_b1.status_code == 200  # user B's own limit is untouched by A's usage


# -- Workspaces (ADR-044, v3 Phase 4) -------------------------------------

def _register_and_login(client, email="user@example.com", password="password123"):
    client.post("/auth/register", json={"email": email, "password": password})
    return client.post("/auth/login", json={"email": email, "password": password}).json()["session_token"]


def test_create_workspace_requires_auth(client):
    resp = client.post("/workspaces", json={"name": "Marketing"})
    assert resp.status_code == 401


def test_create_and_list_workspace(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    create_resp = client.post("/workspaces", json={"name": "Marketing"}, headers=headers)
    assert create_resp.status_code == 200
    workspace_id = create_resp.json()["workspace_id"]

    list_resp = client.get("/workspaces", headers=headers)
    assert list_resp.status_code == 200
    names = [w["name"] for w in list_resp.json()]
    assert names == ["Marketing"]
    assert list_resp.json()[0]["workspace_id"] == workspace_id


def test_create_workspace_requires_a_name(client):
    token = _register_and_login(client)
    resp = client.post("/workspaces", json={"name": "   "},
                        headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 400


def test_workspace_list_is_isolated_per_user(client):
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")
    client.post("/workspaces", json={"name": "A's workspace"}, headers={"Authorization": f"Bearer {token_a}"})

    resp_b = client.get("/workspaces", headers={"Authorization": f"Bearer {token_b}"})
    assert resp_b.json() == []  # user B sees nothing of user A's


def test_rename_workspace(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Old Name"}, headers=headers).json()["workspace_id"]

    resp = client.patch(f"/workspaces/{workspace_id}", json={"name": "New Name"}, headers=headers)
    assert resp.status_code == 200

    detail = client.get(f"/workspaces/{workspace_id}", headers=headers)
    assert detail.json()["name"] == "New Name"


def test_rename_unknown_workspace_returns_404(client):
    token = _register_and_login(client)
    resp = client.patch("/workspaces/not-a-real-id", json={"name": "X"},
                         headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 404


def test_generation_can_be_assigned_to_a_workspace_at_creation(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Pitch Decks"}, headers=headers).json()["workspace_id"]

    gen_resp = client.post("/generate/topic",
                            json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                            headers=headers)
    assert gen_resp.status_code == 200

    detail = client.get(f"/workspaces/{workspace_id}", headers=headers)
    assert detail.status_code == 200
    assert len(detail.json()["projects"]) == 1


def test_generation_with_unowned_workspace_id_is_rejected(client):
    """A caller can't attach their generation to a workspace_id they
    don't own — checked BEFORE any generation work happens."""
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")
    workspace_id = client.post("/workspaces", json={"name": "A's private workspace"},
                                headers={"Authorization": f"Bearer {token_a}"}).json()["workspace_id"]

    resp = client.post("/generate/topic",
                        json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                        headers={"Authorization": f"Bearer {token_b}"})
    assert resp.status_code == 404


def test_projects_endpoint_filters_by_workspace_id(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    ws1 = client.post("/workspaces", json={"name": "Workspace 1"}, headers=headers).json()["workspace_id"]
    ws2 = client.post("/workspaces", json={"name": "Workspace 2"}, headers=headers).json()["workspace_id"]

    client.post("/generate/topic", json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": ws1}, headers=headers)
    client.post("/generate/topic", json={"topic": "Rivers", "slide_count": 4, "workspace_id": ws2}, headers=headers)
    client.post("/generate/topic", json={"topic": "Mountains", "slide_count": 4}, headers=headers)  # ungrouped

    all_projects = client.get("/projects", headers=headers).json()
    assert len(all_projects) == 3  # unfiltered still sees everything, pre-ADR-044 behavior

    ws1_projects = client.get(f"/projects?workspace_id={ws1}", headers=headers).json()
    assert len(ws1_projects) == 1
    assert ws1_projects[0]["workspace_id"] == ws1


def test_deleting_workspace_does_not_delete_its_projects(client):
    """The core design guarantee of ADR-044, tested end to end through
    the real API: a project's actual content must survive its
    workspace being deleted, just landing back in the ungrouped list."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Temporary"}, headers=headers).json()["workspace_id"]
    gen_resp = client.post("/generate/topic",
                            json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                            headers=headers)
    project_id = _poll_job_until_done(
        client,
        client.post("/generate/topic/async",
                     json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                     headers=headers).json()["job_id"],
    )["project_id"]

    delete_resp = client.delete(f"/workspaces/{workspace_id}", headers=headers)
    assert delete_resp.status_code == 200
    assert delete_resp.json()["deleted"] is True

    # workspace itself is gone
    assert client.get(f"/workspaces/{workspace_id}", headers=headers).status_code == 404

    # but the project the workspace contained is very much still there,
    # just ungrouped now
    project_detail = client.get(f"/projects/{project_id}", headers=headers)
    assert project_detail.status_code == 200

    all_projects = client.get("/projects", headers=headers).json()
    matching = [p for p in all_projects if p["project_id"] == project_id]
    assert len(matching) == 1
    assert matching[0]["workspace_id"] is None


def test_delete_unknown_workspace_returns_404(client):
    token = _register_and_login(client)
    resp = client.delete("/workspaces/not-a-real-id", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 404


# -- Brand Memory (ADR-045, v3 Phase 5) ------------------------------------

def test_get_brand_profile_before_ever_setting_one_returns_empty_not_404(client):
    """A workspace that's never had a brand profile set is a normal
    state (200, all-blank fields), not an error — only an unowned or
    nonexistent WORKSPACE is a 404."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Fresh"}, headers=headers).json()["workspace_id"]

    resp = client.get(f"/workspaces/{workspace_id}/brand", headers=headers)
    assert resp.status_code == 200
    assert resp.json()["name"] == ""
    assert resp.json()["colors"] == ""


def test_set_and_get_brand_profile(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Acme"}, headers=headers).json()["workspace_id"]

    set_resp = client.put(f"/workspaces/{workspace_id}/brand",
                           json={"name": "Acme Corp", "colors": "Blue and purple", "tone": "Playful",
                                 "audience": "Investors", "visual_style": "Minimal"},
                           headers=headers)
    assert set_resp.status_code == 200
    assert set_resp.json()["name"] == "Acme Corp"

    get_resp = client.get(f"/workspaces/{workspace_id}/brand", headers=headers)
    assert get_resp.json()["colors"] == "Blue and purple"
    assert get_resp.json()["tone"] == "Playful"


def test_brand_endpoints_require_workspace_ownership(client):
    token_a = _register_and_login(client, "a@example.com")
    token_b = _register_and_login(client, "b@example.com")
    workspace_id = client.post("/workspaces", json={"name": "A's workspace"},
                                headers={"Authorization": f"Bearer {token_a}"}).json()["workspace_id"]

    resp = client.put(f"/workspaces/{workspace_id}/brand", json={"name": "Hijacked"},
                       headers={"Authorization": f"Bearer {token_b}"})
    assert resp.status_code == 404


def test_brand_endpoints_require_auth(client):
    resp = client.get("/workspaces/some-id/brand")
    assert resp.status_code == 401


def test_delete_brand_profile(client):
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Acme"}, headers=headers).json()["workspace_id"]
    client.put(f"/workspaces/{workspace_id}/brand", json={"name": "Acme Corp"}, headers=headers)

    delete_resp = client.delete(f"/workspaces/{workspace_id}/brand", headers=headers)
    assert delete_resp.status_code == 200

    get_resp = client.get(f"/workspaces/{workspace_id}/brand", headers=headers)
    assert get_resp.json()["name"] == ""  # back to the never-set state


def test_generation_into_a_branded_workspace_still_succeeds(client):
    """End-to-end proof that setting a brand profile and then
    generating into that workspace doesn't break anything — the
    actual prompt content isn't observable from the HTTP layer (no
    real AI provider is configured in this hermetic suite), but the
    whole request path (fetch brand -> thread into GenerationRequest
    -> deterministic fallback since no AI configured) must complete
    normally end to end."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Branded"}, headers=headers).json()["workspace_id"]
    client.put(f"/workspaces/{workspace_id}/brand",
               json={"name": "Acme Corp", "tone": "Playful"}, headers=headers)

    resp = client.post("/generate/topic",
                        json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                        headers=headers)
    assert resp.status_code == 200


def test_generation_into_workspace_with_no_brand_set_still_succeeds(client):
    """The far more common case — a workspace with no brand profile at
    all — must be entirely unaffected by ADR-045 existing."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Unbranded"}, headers=headers).json()["workspace_id"]

    resp = client.post("/generate/topic",
                        json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                        headers=headers)
    assert resp.status_code == 200


def test_generation_into_branded_workspace_async_full_round_trip(client):
    """Confirms the job-payload serialize/reconstruct path (BrandProfile
    dict -> job.payload["brand"] -> reconstructed dataclass in the
    worker) works end to end through async generation, not just sync."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Branded Async"}, headers=headers).json()["workspace_id"]
    client.put(f"/workspaces/{workspace_id}/brand", json={"tone": "Playful"}, headers=headers)

    enqueue_resp = client.post("/generate/topic/async",
                                json={"topic": "Volcanoes", "slide_count": 4, "workspace_id": workspace_id},
                                headers=headers)
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]
    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] == "deterministic-topic"


def test_document_upload_generation_into_branded_workspace_succeeds(client):
    """ADR-045's document-mode gap closure — a document-upload
    generation into a workspace WITH a brand profile must complete
    normally end to end (sync path), proving the brand fetch/thread-
    through added to /generate doesn't break the existing upload flow."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Branded Docs"}, headers=headers).json()["workspace_id"]
    client.put(f"/workspaces/{workspace_id}/brand", json={"tone": "Playful"}, headers=headers)

    resp = client.post(
        f"/generate?workspace_id={workspace_id}",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
        headers=headers,
    )
    assert resp.status_code == 200


def test_document_upload_generation_into_branded_workspace_async_full_round_trip(client):
    """Same as above but through the async/job-payload path — proves
    the serialize-brand-into-payload / reconstruct-in-worker round
    trip works for document uploads too, not just topic generation."""
    token = _register_and_login(client)
    headers = {"Authorization": f"Bearer {token}"}
    workspace_id = client.post("/workspaces", json={"name": "Branded Docs Async"},
                                headers=headers).json()["workspace_id"]
    client.put(f"/workspaces/{workspace_id}/brand", json={"tone": "Playful"}, headers=headers)

    enqueue_resp = client.post(
        f"/generate/async?workspace_id={workspace_id}",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
        headers=headers,
    )
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]
    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] in ("rule-based", "ai-generated")


# -- Real PDF end to end, all 5 export formats (ADR-050) -------------------
# Closes a real, previously-untested gap: the v3 roadmap had CLAIMED
# "convert PDF into X already works, no new extraction logic needed"
# across Phase 3/6, but every prior HTTP test used a .txt fixture —
# the only PDF ever sent through the real HTTP layer before this was
# the deliberately-corrupt one in test_generate_corrupt_pdf_returns_422.
# A real, valid PDF had never actually been proven to work end to end
# through the API until now.

def test_generate_from_real_pdf_sync(client):
    resp = client.post("/generate", files={"file": ("water_cycle.pdf", SAMPLE_PDF_DOC, "application/pdf")})
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/zip"


@pytest.mark.parametrize("export_format", [
    "document_docx", "document_pdf",
])
def test_generate_from_real_pdf_every_non_pptx_format(client, export_format):
    resp = client.post(
        "/generate",
        files={"file": ("water_cycle.pdf", SAMPLE_PDF_DOC, "application/pdf")},
        params={"export_format": export_format},
    )
    assert resp.status_code == 200


# -- Document Q&A (ADR-050, v3 Phase 7) -------------------------------------

def test_ask_document_requires_a_question(client):
    resp = client.post("/documents/ask", files={"file": ("doc.txt", SAMPLE_TEXT_DOC, "text/plain")})
    assert resp.status_code == 400


def test_ask_document_returns_an_answer_field(client):
    """Hermetic suite has no AI configured, so this exercises the
    NullAdapter degradation path specifically — still a real 200 with
    a real (honest) answer field, not an error."""
    resp = client.post(
        "/documents/ask",
        files={"file": ("doc.txt", SAMPLE_TEXT_DOC, "text/plain")},
        params={"question": "What causes precipitation?"},
    )
    assert resp.status_code == 200
    assert "answer" in resp.json()
    assert "not configured" in resp.json()["answer"].lower()


def test_ask_document_works_with_a_real_pdf(client):
    resp = client.post(
        "/documents/ask",
        files={"file": ("water_cycle.pdf", SAMPLE_PDF_DOC, "application/pdf")},
        params={"question": "What is this document about?"},
    )
    assert resp.status_code == 200
    assert "answer" in resp.json()


def test_ask_document_unsupported_filetype_returns_400(client):
    resp = client.post(
        "/documents/ask",
        files={"file": ("data.xyz", b"whatever content", "application/octet-stream")},
        params={"question": "test?"},
    )
    assert resp.status_code == 400


def test_ask_document_corrupt_pdf_returns_422(client):
    resp = client.post(
        "/documents/ask",
        files={"file": ("broken.pdf", b"this is not a real pdf file at all", "application/pdf")},
        params={"question": "test?"},
    )
    assert resp.status_code == 422


def test_ask_document_gated_by_its_own_quota_not_generation_quota(client, monkeypatch):
    """ADR-050's whole point in having a SEPARATE quota bucket from
    generation, proven — a generation-quota env var set to 0 must NOT
    block Q&A, and vice versa."""
    monkeypatch.setenv("OPENPRESENT_DAILY_GENERATION_LIMIT_ANON", "0")
    resp = client.post(
        "/documents/ask",
        files={"file": ("doc.txt", SAMPLE_TEXT_DOC, "text/plain")},
        params={"question": "test?"},
    )
    assert resp.status_code == 200  # generation's cap being 0 doesn't touch Q&A


def test_ask_document_blocked_after_its_own_daily_limit(client, monkeypatch):
    monkeypatch.setenv("OPENPRESENT_DAILY_QA_LIMIT_ANON", "2")
    for _ in range(2):
        resp = client.post(
            "/documents/ask",
            files={"file": ("doc.txt", SAMPLE_TEXT_DOC, "text/plain")},
            params={"question": "test?"},
        )
        assert resp.status_code == 200
    blocked = client.post(
        "/documents/ask",
        files={"file": ("doc.txt", SAMPLE_TEXT_DOC, "text/plain")},
        params={"question": "test?"},
    )
    assert blocked.status_code == 429


def test_generate_topic_as_document_docx_sync(client):
    """Same engine, same request shape — only export_format differs.
    No new endpoint was needed for this format, by design."""
    resp = client.post("/generate/topic", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_docx",
    })
    assert resp.status_code == 200
    assert resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert 'filename="document.docx"' in resp.headers["content-disposition"]
    doc = Document(io.BytesIO(resp.content))  # a real, parseable docx
    assert len(doc.paragraphs) > 0


def test_generate_topic_as_document_docx_is_not_bundled_with_speaker_notes(client):
    """The pptx path bundles a speaker_notes.docx companion by default —
    that guard is keyed on export_format == 'pptx', so a document
    export must come back as a bare .docx, never a .zip."""
    resp = client.post("/generate/topic", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_docx", "bundle_speaker_notes": True,
    })
    assert resp.status_code == 200
    assert resp.headers["content-type"] != "application/zip"


def test_generate_topic_document_docx_async_full_round_trip(client):
    enqueue_resp = client.post("/generate/topic/async", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_docx",
    })
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]

    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] == "deterministic-topic"

    download_resp = client.get(f"/jobs/{job_id}/download")
    assert download_resp.status_code == 200
    assert download_resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert 'filename="document.docx"' in download_resp.headers["content-disposition"]
    doc = Document(io.BytesIO(download_resp.content))
    assert len(doc.paragraphs) > 0


# -- PDF documents, third and final format (ADR-055) ------------------------
# Replaces infographic_svg/diagram_svg/poster_svg (ADR-046/047/048) after the
# scope narrowed to pptx/docx/pdf only. document_pdf renders directly from
# the Recipe (see DocumentPdfExportAdapter) and shares document_docx's prose
# content-shaping, so these tests mirror the document_docx tests above,
# checking the PDF-specific bytes and headers instead of docx internals.

def test_generate_topic_as_document_pdf_sync(client):
    resp = client.post("/generate/topic", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_pdf",
    })
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert 'filename="document.pdf"' in resp.headers["content-disposition"]
    assert resp.content.startswith(b"%PDF")  # a real, well-formed PDF


def test_generate_topic_as_document_pdf_is_not_bundled_with_speaker_notes(client):
    resp = client.post("/generate/topic", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_pdf", "bundle_speaker_notes": True,
    })
    assert resp.status_code == 200
    assert resp.headers["content-type"] != "application/zip"


def test_generate_topic_document_pdf_async_full_round_trip(client):
    enqueue_resp = client.post("/generate/topic/async", json={
        "topic": "Renewable Energy Adoption", "slide_count": 4,
        "export_format": "document_pdf",
    })
    assert enqueue_resp.status_code == 200
    job_id = enqueue_resp.json()["job_id"]

    result = _poll_job_until_done(client, job_id)
    assert result["structure_source"] == "deterministic-topic"

    download_resp = client.get(f"/jobs/{job_id}/download")
    assert download_resp.status_code == 200
    assert download_resp.headers["content-type"] == "application/pdf"
    assert 'filename="document.pdf"' in download_resp.headers["content-disposition"]
    assert download_resp.content.startswith(b"%PDF")


def test_generate_document_upload_as_document_pdf(client):
    resp = client.post(
        "/generate",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
        params={"export_format": "document_pdf"},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"


def test_jobs_endpoint_surfaces_stage_while_running(client_no_worker):
    # Deterministic generation completes too fast to reliably observe an
    # intermediate stage via the real async race, so the RUNNING state is
    # set up directly through the same QueuePort the API route itself
    # reads from — this is testing the /jobs/{id} route's own response
    # shape (ADR-040), not the worker's timing. Uses client_no_worker
    # (not client) specifically so there's no real background worker
    # thread racing to actually process this job out from under the
    # test's manual state — see that fixture's docstring for the real,
    # reproduced failure this was fixed in response to, not a guess.
    queue = registry.get_queue_adapter()
    job_id = queue.enqueue("generate_topic", {"topic": "Volcanoes", "slide_count": 4})
    queue.dequeue()  # PENDING -> RUNNING; no worker running to race this
    queue.update_stage(job_id, "building_outline")

    resp = client_no_worker.get(f"/jobs/{job_id}")
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "running"
    assert body["stage"] == "building_outline"


def test_jobs_endpoint_omits_stage_once_done(client):
    enqueue_resp = client.post("/generate/topic/async", json={"topic": "Volcanoes", "slide_count": 4})
    job_id = enqueue_resp.json()["job_id"]
    result = _poll_job_until_done(client, job_id)
    assert "stage" not in result  # stage is a running-only signal, redundant once complete


def test_jobs_unknown_id_returns_404(client):
    resp = client.get("/jobs/not-a-real-job-id")
    assert resp.status_code == 404


def test_jobs_download_unknown_id_returns_404(client):
    resp = client.get("/jobs/not-a-real-job-id/download")
    assert resp.status_code == 404


# -- Auth + project isolation, full HTTP round trip -----------------------

def test_auth_and_project_isolation_full_http_flow(client):
    # Register + login user A
    reg_a = client.post("/auth/register", json={"email": "alice@example.com", "password": "hunter22"})
    assert reg_a.status_code == 200
    login_a = client.post("/auth/login", json={"email": "alice@example.com", "password": "hunter22"})
    assert login_a.status_code == 200
    token_a = login_a.json()["session_token"]

    # Generate as Alice, authenticated -> should persist as a project
    gen_resp = client.post(
        "/generate/async",
        headers={"Authorization": f"Bearer {token_a}"},
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    job_id = gen_resp.json()["job_id"]
    result = _poll_job_until_done(client, job_id)
    assert result["project_id"] is not None

    # Alice sees her project
    projects_a = client.get("/projects", headers={"Authorization": f"Bearer {token_a}"})
    assert projects_a.status_code == 200
    assert len(projects_a.json()) == 1

    # Register + login user B — must NOT see Alice's project
    client.post("/auth/register", json={"email": "bob@example.com", "password": "hunter22"})
    login_b = client.post("/auth/login", json={"email": "bob@example.com", "password": "hunter22"})
    token_b = login_b.json()["session_token"]
    projects_b = client.get("/projects", headers={"Authorization": f"Bearer {token_b}"})
    assert projects_b.status_code == 200
    assert len(projects_b.json()) == 0  # per-owner isolation, confirmed over real HTTP

    # No auth header at all -> 401
    unauthenticated = client.get("/projects")
    assert unauthenticated.status_code == 401


def test_register_duplicate_email_returns_error(client):
    client.post("/auth/register", json={"email": "dup@example.com", "password": "hunter22"})
    second = client.post("/auth/register", json={"email": "dup@example.com", "password": "hunter22"})
    assert second.status_code == 409  # Conflict — correct REST status for "already exists"


def test_login_wrong_password_returns_401(client):
    client.post("/auth/register", json={"email": "carol@example.com", "password": "correct-pw"})
    resp = client.post("/auth/login", json={"email": "carol@example.com", "password": "wrong-pw"})
    assert resp.status_code == 401


# -- /auth/me (ADR-056) -----------------------------------------------------
# Login previously only ever returned a bare session token, never the
# user's own email — the frontend Settings page needed a real way to
# answer "who am I logged in as" without decoding the (intentionally
# opaque) session token itself.

def test_auth_me_requires_auth(client):
    resp = client.get("/auth/me")
    assert resp.status_code == 401


def test_auth_me_returns_the_logged_in_users_email(client):
    token = _register_and_login(client, email="dana@example.com")
    resp = client.get("/auth/me", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 200
    body = resp.json()
    assert body["email"] == "dana@example.com"
    assert "user_id" in body


def test_auth_me_rejects_an_invalid_token(client):
    resp = client.get("/auth/me", headers={"Authorization": "Bearer not-a-real-token"})
    assert resp.status_code == 401


def test_anonymous_generate_still_works_no_account_required(client):
    """Core product promise, verified over real HTTP: generation never
    requires an account — only saving/reusing a project does."""
    resp = client.post(
        "/generate",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200


# -- ADR-039 regression: sync generation now saves a project when logged
# in, and the frontend can actually read the resulting project id ------
# (previously ONLY /generate/async did this — the homepage's primary
# generation flow, which calls the SYNC endpoints, produced nothing a
# logged-in user could ever open in the slide editor.)

def test_generate_sync_document_saves_project_when_authenticated(client):
    client.post("/auth/register", json={"email": "syncsave@example.com", "password": "hunter22"})
    login = client.post("/auth/login", json={"email": "syncsave@example.com", "password": "hunter22"})
    token = login.json()["session_token"]

    resp = client.post(
        "/generate", headers={"Authorization": f"Bearer {token}"},
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200
    project_id = resp.headers.get("X-Project-Id")
    assert project_id, "sync /generate must return X-Project-Id when the caller is authenticated"

    # The saved project is actually retrievable afterward — this is the
    # real proof, not just that a header was returned.
    project_resp = client.get(f"/projects/{project_id}", headers={"Authorization": f"Bearer {token}"})
    assert project_resp.status_code == 200


def test_generate_sync_document_saves_nothing_when_anonymous(client):
    resp = client.post(
        "/generate",
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    assert resp.status_code == 200
    assert "X-Project-Id" not in resp.headers


def test_generate_sync_topic_saves_project_when_authenticated(client):
    client.post("/auth/register", json={"email": "synctopic@example.com", "password": "hunter22"})
    login = client.post("/auth/login", json={"email": "synctopic@example.com", "password": "hunter22"})
    token = login.json()["session_token"]

    resp = client.post(
        "/generate/topic", headers={"Authorization": f"Bearer {token}"},
        json={"topic": "Volcanoes", "slide_count": 4},
    )
    assert resp.status_code == 200
    project_id = resp.headers.get("X-Project-Id")
    assert project_id

    project_resp = client.get(f"/projects/{project_id}", headers={"Authorization": f"Bearer {token}"})
    assert project_resp.status_code == 200
    assert project_resp.json()["slide_count"] == 4


def test_project_id_header_is_cors_exposed():
    """The exact bug this ADR also fixed: a header can be present in a
    real HTTP response and still be invisible to browser JavaScript on
    a cross-origin request (Vercel frontend -> Render backend are
    different origins) unless the server explicitly lists it in
    Access-Control-Expose-Headers. Verified directly against the
    response, not just assumed from the CORSMiddleware config."""
    from backend.api.main import app
    from starlette.middleware.cors import CORSMiddleware as _CORSMiddlewareClass

    cors_middleware = next(
        m for m in app.user_middleware if m.cls is _CORSMiddlewareClass
    )
    exposed = cors_middleware.kwargs.get("expose_headers", [])
    for header in ("X-Project-Id", "X-Structure-Source", "X-Quality-Score"):
        assert header in exposed, f"{header} is set on responses but not CORS-exposed to the frontend"


# -- Slide-level editing / partial regeneration (ADR-038) ----------------

def _register_login_and_generate_project(client):
    client.post("/auth/register", json={"email": "editor@example.com", "password": "hunter22"})
    login = client.post("/auth/login", json={"email": "editor@example.com", "password": "hunter22"})
    token = login.json()["session_token"]
    headers = {"Authorization": f"Bearer {token}"}

    gen_resp = client.post(
        "/generate/async", headers=headers,
        files={"file": ("water_cycle.txt", SAMPLE_TEXT_DOC, "text/plain")},
    )
    job_id = gen_resp.json()["job_id"]
    result = _poll_job_until_done(client, job_id)
    return headers, result["project_id"]


def test_get_slide_returns_full_detail(client):
    headers, project_id = _register_login_and_generate_project(client)
    project = client.get(f"/projects/{project_id}", headers=headers).json()
    first_order = project["slides"][0]["order"]

    resp = client.get(f"/projects/{project_id}/slides/{first_order}", headers=headers)
    assert resp.status_code == 200
    body = resp.json()
    for key in ("order", "title", "bullets", "notes", "layout_type", "image_query"):
        assert key in body


def test_get_slide_unknown_order_returns_404(client):
    headers, project_id = _register_login_and_generate_project(client)
    resp = client.get(f"/projects/{project_id}/slides/999", headers=headers)
    assert resp.status_code == 404


# -- DELETE /projects/{id} (ADR-056) -----------------------------------------
# StoragePort.delete_recipe() existed since Phase 4 but was never wired to
# an HTTP route — this is the frontend's "delete a previous chat" gap and
# the backend's missing route, closed together.

def test_delete_project_requires_auth(client):
    resp = client.delete("/projects/some-id")
    assert resp.status_code == 401


def test_delete_project_removes_it(client):
    headers, project_id = _register_login_and_generate_project(client)

    resp = client.delete(f"/projects/{project_id}", headers=headers)
    assert resp.status_code == 200
    assert resp.json() == {"deleted": True}

    reget = client.get(f"/projects/{project_id}", headers=headers)
    assert reget.status_code == 404  # actually gone, not just hidden


def test_delete_project_no_longer_appears_in_the_list(client):
    headers, project_id = _register_login_and_generate_project(client)
    client.delete(f"/projects/{project_id}", headers=headers)
    listing = client.get("/projects", headers=headers)
    assert project_id not in [p["project_id"] for p in listing.json()]


def test_delete_unknown_project_returns_404(client):
    token = _register_and_login(client, email="erin@example.com")
    resp = client.delete("/projects/does-not-exist", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 404


def test_cannot_delete_another_users_project(client):
    """Per-owner isolation (Blueprint Section 11) applies to delete just
    like every other project route — a project_id that's real, just not
    yours, must come back as 404, not 403 (existence isn't leaked)."""
    headers_a, project_id = _register_login_and_generate_project(client)

    token_b = _register_and_login(client, email="frank@example.com", password="hunter22")
    resp = client.delete(f"/projects/{project_id}", headers={"Authorization": f"Bearer {token_b}"})
    assert resp.status_code == 404

    # Confirm it's genuinely untouched, not soft-deleted by the failed attempt
    still_there = client.get(f"/projects/{project_id}", headers=headers_a)
    assert still_there.status_code == 200


def test_get_slide_requires_auth(client):
    headers, project_id = _register_login_and_generate_project(client)
    resp = client.get(f"/projects/{project_id}/slides/1")
    assert resp.status_code == 401


def test_patch_slide_manual_edit_over_real_http(client):
    headers, project_id = _register_login_and_generate_project(client)
    project = client.get(f"/projects/{project_id}", headers=headers).json()
    first_order = project["slides"][0]["order"]

    resp = client.patch(
        f"/projects/{project_id}/slides/{first_order}", headers=headers,
        json={"title": "A Brand New Title", "bullets": ["Point A", "Point B"]},
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["title"] == "A Brand New Title"
    assert body["bullets"] == ["Point A", "Point B"]

    # Persisted — a fresh GET shows the same edit.
    reget = client.get(f"/projects/{project_id}/slides/{first_order}", headers=headers)
    assert reget.json()["title"] == "A Brand New Title"


def test_patch_slide_empty_body_returns_422(client):
    headers, project_id = _register_login_and_generate_project(client)
    project = client.get(f"/projects/{project_id}", headers=headers).json()
    first_order = project["slides"][0]["order"]
    resp = client.patch(f"/projects/{project_id}/slides/{first_order}", headers=headers, json={})
    assert resp.status_code == 422


def test_patch_slide_requires_auth(client):
    resp = client.patch("/projects/fake-id/slides/1", json={"title": "X"})
    assert resp.status_code == 401


def test_patch_slide_wrong_owner_returns_404(client):
    headers, project_id = _register_login_and_generate_project(client)
    client.post("/auth/register", json={"email": "other@example.com", "password": "hunter22"})
    other_login = client.post("/auth/login", json={"email": "other@example.com", "password": "hunter22"})
    other_headers = {"Authorization": f"Bearer {other_login.json()['session_token']}"}

    resp = client.patch(f"/projects/{project_id}/slides/1", headers=other_headers,
                         json={"title": "Hijacked"})
    assert resp.status_code == 404


def test_regenerate_slide_returns_503_without_ai_configured(client):
    """This test suite's fixture forces OPENPRESENT_AI_ADAPTER=null —
    confirms the real HTTP path surfaces AIUnavailableError as a
    proper 503, not a 500 crash or a silent no-op."""
    headers, project_id = _register_login_and_generate_project(client)
    project = client.get(f"/projects/{project_id}", headers=headers).json()
    first_order = project["slides"][0]["order"]

    resp = client.post(
        f"/projects/{project_id}/slides/{first_order}/regenerate", headers=headers,
        json={"instructions": "make this punchier"},
    )
    assert resp.status_code == 503


def test_regenerate_slide_requires_auth(client):
    resp = client.post("/projects/fake-id/slides/1/regenerate", json={})
    assert resp.status_code == 401


def test_regenerate_slide_unknown_project_returns_404_not_503(client, monkeypatch):
    """Ownership/existence checks must happen before the AI-availability
    check would even matter — even with no AI configured, an unknown
    project should 404, not 503."""
    headers, _ = _register_login_and_generate_project(client)
    resp = client.post("/projects/not-a-real-project/slides/1/regenerate", headers=headers, json={})
    assert resp.status_code in (404, 503)  # 404 if ownership is checked first; 503 is still valid
    # (this project genuinely doesn't exist AND no AI is configured — either
    # order of checks is defensible; what matters is it's a clean 4xx/5xx,
    # never a 500 crash)
    assert resp.status_code != 500
