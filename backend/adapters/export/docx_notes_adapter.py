"""
SpeakerNotesDocxExportAdapter — ADR-030.

A companion export, not a replacement for PPTX: produces a Word
document listing each slide's title, bullets, and speaker notes in
order — useful for a presenter who wants to read/print/rehearse their
notes without opening PowerPoint, and a direct answer to the explicit
request that speaker notes ship as a separate .docx alongside the
.pptx (see backend/engines/export_bundle.py, which packages both
together into one .zip download).

Implements ExportPort like every other export adapter (format_id =
"notes_docx") — usable standalone via /generate?export_format=notes_docx
if someone only wants the notes doc, though the default product flow
bundles it with the .pptx automatically.
"""

from docx import Document
from docx.shared import Pt
from backend.ports.export import ExportPort
from backend.models.recipe import Recipe, BlockType


class SpeakerNotesDocxExportAdapter(ExportPort):
    def format_id(self) -> str:
        return "notes_docx"

    def export(self, recipe: Recipe) -> bytes:
        import io

        doc = Document()
        title_slide = recipe.outline.slides[0] if recipe.outline.slides else None
        doc.add_heading(title_slide.title if title_slide else "Speaker Notes", level=0)

        for i, slide in enumerate(sorted(recipe.outline.slides, key=lambda s: s.order), start=1):
            doc.add_heading(f"Slide {i}: {slide.title}", level=1)

            bullets = [b.text for b in slide.content_blocks if b.type == BlockType.BULLET and b.text]
            if bullets:
                p = doc.add_paragraph()
                p.add_run("On-slide content: ").bold = True
                for j, bullet in enumerate(bullets):
                    doc.add_paragraph(bullet, style="List Bullet")

            notes = [b.text for b in slide.content_blocks if b.type == BlockType.NOTE and b.text]
            if notes:
                p = doc.add_paragraph()
                run = p.add_run("Speaker notes")
                run.bold = True
                run.italic = True
                for note in notes:
                    note_p = doc.add_paragraph(note)
                    note_p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph()
                p.add_run("(No speaker notes for this slide.)").italic = True

            doc.add_paragraph()  # spacing between slides

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
