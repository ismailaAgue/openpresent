"""
DocumentDocxExportAdapter — ADR-041 (v3 Phase 3, "Documents" output type).
Design/typography reworked in ADR-054.

Renders the exact same Recipe/Outline the slide pipeline produces, but
as a flowing Word document instead of a deck: each Slide becomes a
document section (its title -> a heading, its bullets -> either real
prose paragraphs or a real bullet list, depending on how the content
was actually written — see _all_read_as_prose), rather than being laid
out for on-screen display.

This is deliberately NOT the same thing as SpeakerNotesDocxExportAdapter
(docx_notes_adapter.py) — that one is a presenter's companion aid
("Slide 3: Market Opportunity" headers, "On-slide content:" labels).
This adapter is meant to BE the deliverable itself: a real report,
proposal, or exec summary a reader opens and reads top to bottom, with
no reference to slides anywhere in it. Per the v3 roadmap's Phase 3:
reuse Strategy/Outline/Content generation as-is (same Outline model,
same generate_presentation_from_topic()) and only Layout/Export differ
— this file is that "Export differs" half. There is no separate
document-generation pipeline; export_format="document_docx" is a
normal ExportPort choice on the same engine call as every other format.

ADR-054 — two real, user-reported problems fixed together, both found
by actually rendering output and looking at it (docx skill's own
guidance), not by trusting a green test suite:
1. Content used to always be terse slide-bullet fragments regardless
   of format (fixed upstream — json_pipeline_base.py's content prompt
   and rule_based.py's text chunker are now format-aware). This
   adapter's job is just to render whatever shape the content genuinely
   is: real paragraphs stay real paragraphs, real lists stay real
   lists, per _all_read_as_prose.
2. This file never used the project's theme at all — every document
   rendered in Word's default blue heading color regardless of the
   chosen theme, unlike every other export format (pptx, infographic,
   diagram, poster), which all consistently apply _COLOR_SETS. The
   title page also relied on python-docx's built-in "Title" style,
   whose bottom-border box doesn't span the actual centered text width
   — rendering as visually left-anchored even though the text itself
   was centered (confirmed by rendering an isolated test file and
   looking, not guessed at). Both fixed: real theme colors applied to
   the title and every heading, and a hand-built title block using
   direct paragraph formatting plus a full-page-width accent rule
   (a paragraph bottom border, per this project's docx skill: "don't
   use a table as a horizontal rule").
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from backend.ports.export import ExportPort
from backend.models.recipe import Recipe, BlockType
from backend.adapters.export.pptx_adapter import _COLOR_SETS


def _rgb(color: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*color)


def _set_bottom_border(paragraph, rgb_hex: str, size: int = 18) -> None:
    """Adds a full-page-width bottom border to a paragraph, used as a
    horizontal rule. A paragraph border spans between the page's actual
    margins regardless of its content's width — unlike python-docx's
    built-in "Title" style, whose border box is sized to the style's
    own fixed width, not the page or the centered text (the bug this
    replaces, see module docstring). python-docx has no high-level API
    for paragraph borders, so this sets the OOXML directly — a standard,
    well-known technique for this specific gap in the library."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


class DocumentDocxExportAdapter(ExportPort):
    def format_id(self) -> str:
        return "document_docx"

    def export(self, recipe: Recipe) -> bytes:
        import io

        colors = _COLOR_SETS.get(recipe.theme.color_set_id, _COLOR_SETS["neutral"])
        title_color = _rgb(colors["title"])
        accent_color = _rgb(colors["accent"])
        accent_hex = "{:02X}{:02X}{:02X}".format(*colors["accent"])

        slides = sorted(recipe.outline.slides, key=lambda s: s.order)
        doc = Document()

        self._add_title_page(doc, slides[0].title if slides else "Untitled Document",
                              title_color, accent_hex)

        # slides[0] became the title page above; everything else is a
        # real section. A deck's closing slide (e.g. "Questions?",
        # "Contact Information") reads fine as a document's final
        # section too — no special-casing needed there.
        for slide in slides[1:]:
            heading = doc.add_heading(slide.title, level=1)
            for run in heading.runs:
                run.font.color.rgb = title_color

            bullets = [b.text for b in slide.content_blocks if b.type == BlockType.BULLET and b.text]
            if self._all_read_as_prose(bullets):
                # ADR-054: every bullet here is itself a complete
                # paragraph (the AI content prompt, for this format,
                # asks for 1-3 short paragraphs as separate entries —
                # see json_pipeline_base.py's format branch), so each
                # one becomes its own real paragraph. The original
                # single-bullet case is just the n=1 instance of this
                # same rule, not special-cased separately anymore.
                for paragraph_text in bullets:
                    para = doc.add_paragraph(paragraph_text)
                    para.paragraph_format.space_after = Pt(10)
            else:
                for bullet in bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

            # Notes exist in the model to support a *presenter* reading
            # a deck aloud — irrelevant once the deliverable is the
            # document itself, so intentionally not rendered here
            # (unlike SpeakerNotesDocxExportAdapter, where they're the
            # entire point).

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _all_read_as_prose(bullets: list[str]) -> bool:
        """True if EVERY bullet in this section reads as a complete
        sentence/paragraph (ends in terminal punctuation) — meaning
        the content was authored as prose, not a list, regardless of
        how many separate paragraph-bullets there are. A single short
        fragment (no terminal punctuation) or a genuine multi-item
        list (each item a short phrase, not a sentence) stays a real
        bulleted list; a lone bullet ending in punctuation is the
        n=1 case of this same rule, not a separate special case."""
        return len(bullets) >= 1 and all(b.rstrip().endswith((".", "!", "?")) for b in bullets)

    @staticmethod
    def _add_title_page(doc: Document, title: str, title_color: RGBColor, accent_hex: str) -> None:
        # Direct formatting instead of the built-in "Title" style — see
        # module docstring for exactly why (the built-in style's border
        # box doesn't track the actual centered text width, so it reads
        # as left-anchored once rendered even with alignment=CENTER).
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(180)  # vertically balances the page, not pinned to the very top
        title_para.paragraph_format.space_after = Pt(16)
        run = title_para.add_run(title)
        run.font.size = Pt(30)
        run.font.bold = True
        run.font.color.rgb = title_color

        rule_para = doc.add_paragraph()
        rule_para.paragraph_format.space_after = Pt(0)
        _set_bottom_border(rule_para, accent_hex)

        doc.add_page_break()
